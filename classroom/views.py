from io import BytesIO
import base64
import json
import os
import re
import uuid

from django.conf import settings

from django.shortcuts import render, redirect
from django.contrib import messages
from django.contrib.auth.models import User
from django.contrib.auth import authenticate, login, logout
from django.contrib.auth.decorators import login_required
from django.core.files.base import ContentFile
from django.http import JsonResponse
from django.utils import timezone
from django.utils.dateparse import parse_datetime
from django.views.decorators.http import require_POST

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .models import (
    Profile,
    Classroom,
    JoinRequest,
    Homework,
    HomeworkAttachment,
    StudentHomeworkFile,
    EditorImage,
)


# ─────────────────────────────────────────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────────────────────────────────────────

def build_homework_filename(student, homework, classroom):
    """
    Format: username_homework_title_studentcode_teachername.docx
    Spaces in homework title are replaced with underscores.
    Example: john_hello_world_1234567_mrssmith.docx
    """
    safe_title   = homework.title.replace(' ', '_').lower()
    student_code = student.profile.unique_code
    teacher_name = classroom.teacher.username.lower().replace(' ', '_')
    return f"{student.username}_{safe_title}_{student_code}_{teacher_name}.docx"


def build_docx_bytes(student, classroom, homework, editor_html=None):
    """
    Create the .docx file and return raw bytes.
    Images in editor_html are read directly from MEDIA_ROOT (for uploaded
    /media/... URLs) or decoded from base64 — so all images are correctly
    embedded regardless of how many there are.
    """
    from docx.shared import Inches

    document = Document()

    document.add_heading(homework.title, 0)
    document.add_paragraph(f"Student: {student.username}")
    document.add_paragraph(f"Student code: {student.profile.unique_code}")
    document.add_paragraph(f"Classroom: {classroom.name}")
    document.add_paragraph(f"Teacher: {classroom.teacher.username}")

    if homework.deadline:
        document.add_paragraph(f"Deadline: {homework.deadline}")
    else:
        document.add_paragraph("Deadline: Not specified")

    document.add_paragraph("")
    document.add_paragraph("Homework description:")
    document.add_paragraph(homework.description)

    if homework.attachments.exists():
        document.add_paragraph("")
        document.add_paragraph("Teacher attachments:")
        for attachment in homework.attachments.all():
            desc = attachment.description or "No description"
            document.add_paragraph(f"- {attachment.file.name} | {desc}")

    document.add_paragraph("")
    document.add_paragraph("Student answer:")
    document.add_paragraph("")

    if editor_html:
        # Split HTML on every <img ...> tag so we can handle text and images
        # in document order, supporting any number of images.
        parts = re.split(r'(<img[^>]+?>)', editor_html, flags=re.IGNORECASE | re.DOTALL)

        any_content = False

        for part in parts:
            img_match = re.search(r'<img[^>]+?src=["\']([^"\']+)["\']', part, re.IGNORECASE)

            if img_match:
                src = img_match.group(1)
                img_bytes = None

                try:
                    if src.startswith('data:image'):
                        # ── base64 inline ──────────────────────────────
                        _header, b64data = src.split(',', 1)
                        img_bytes = base64.b64decode(b64data)

                    elif src.startswith(settings.MEDIA_URL):
                        # ── Uploaded file — read directly from MEDIA_ROOT ──
                        # e.g. /media/editor_images/foo.png  →  MEDIA_ROOT/editor_images/foo.png
                        relative_path = src[len(settings.MEDIA_URL):]
                        abs_path = os.path.join(settings.MEDIA_ROOT, relative_path)
                        if os.path.isfile(abs_path):
                            with open(abs_path, 'rb') as f:
                                img_bytes = f.read()

                    # external http(s) URLs: skip silently (can't fetch from worker)

                except Exception:
                    img_bytes = None

                if img_bytes:
                    try:
                        document.add_picture(BytesIO(img_bytes), width=Inches(5))
                        any_content = True
                    except Exception:
                        document.add_paragraph("[Image could not be embedded]")
                else:
                    document.add_paragraph(f"[Image: {src[:80]}]")

            else:
                # Text / other HTML — strip tags, collapse whitespace
                plain = re.sub(r'<[^>]+>', ' ', part)
                plain = re.sub(r'\s+', ' ', plain).strip()
                if plain:
                    document.add_paragraph(plain)
                    any_content = True

        if not any_content:
            document.add_paragraph("(empty)")

    else:
        for _ in range(12):
            document.add_paragraph("____________________________________________________________")

    stream = BytesIO()
    document.save(stream)
    stream.seek(0)
    return stream.read()


def create_student_homework_docx(student, classroom, homework):
    """Create initial blank docx for a student. Returns (filename, ContentFile)."""
    filename  = build_homework_filename(student, homework, classroom)
    raw_bytes = build_docx_bytes(student, classroom, homework)
    return filename, ContentFile(raw_bytes)


def create_missing_student_files_for_homework(homework, classroom):
    accepted_students = JoinRequest.objects.filter(
        classroom=classroom,
        status='accepted'
    ).select_related('student', 'student__profile')

    created_count = 0
    for join in accepted_students:
        student = join.student
        if StudentHomeworkFile.objects.filter(homework=homework, student=student).exists():
            continue

        filename, doc_file = create_student_homework_docx(student, classroom, homework)
        obj = StudentHomeworkFile(homework=homework, classroom=classroom, student=student)
        obj.file.save(filename, doc_file, save=True)
        created_count += 1

    return created_count


def create_missing_files_for_student_in_classroom(student, classroom):
    homeworks = Homework.objects.filter(classroom=classroom).prefetch_related('attachments')
    created_count = 0

    for homework in homeworks:
        if StudentHomeworkFile.objects.filter(homework=homework, student=student).exists():
            continue

        filename, doc_file = create_student_homework_docx(student, classroom, homework)
        obj = StudentHomeworkFile(homework=homework, classroom=classroom, student=student)
        obj.file.save(filename, doc_file, save=True)
        created_count += 1

    return created_count


def _refresh_docx_file(student_file):
    """Regenerate the physical .docx from current editor_content and save it."""
    student  = student_file.student
    homework = student_file.homework
    classroom = student_file.classroom

    filename  = build_homework_filename(student, homework, classroom)
    raw_bytes = build_docx_bytes(student, classroom, homework,
                                  editor_html=student_file.editor_content)

    # Delete old file, save new one with the correct name
    student_file.file.delete(save=False)
    student_file.file.save(filename, ContentFile(raw_bytes), save=False)


# ─────────────────────────────────────────────────────────────────────────────
# AUTH VIEWS
# ─────────────────────────────────────────────────────────────────────────────

def login_and_register_page(request):
    if request.user.is_authenticated:
        return redirect('main_page')
    return render(request, 'auth_for_classroom.html')


def register_to_classroom(request):
    if request.method != 'POST':
        return redirect('login_and_register_page')

    username = request.POST.get('username', '').strip()
    password = request.POST.get('password', '').strip()
    role     = request.POST.get('role', '').strip()

    if not username or not password or not role:
        messages.error(request, "Barcha maydonlarni to'ldiring.")
        return redirect('login_and_register_page')

    if len(password) < 8:
        messages.error(request, "Parol kamida 8 ta belgidan iborat bo'lishi kerak.")
        return redirect('login_and_register_page')

    if role not in ['teacher', 'student']:
        messages.error(request, "Role noto'g'ri tanlangan.")
        return redirect('login_and_register_page')

    if User.objects.filter(username=username).exists():
        messages.error(request, "Bu username allaqachon mavjud.")
        return redirect('login_and_register_page')

    user    = User.objects.create_user(username=username, password=password)
    profile = Profile.objects.create(user=user, role=role)

    login(request, user)
    messages.success(
        request,
        f"Account yaratildi. Username: {user.username}, Role: {profile.role}, Code: {profile.unique_code}"
    )
    return redirect('main_page')


def login_to_classroom(request):
    if request.method != 'POST':
        return redirect('login_and_register_page')

    username = request.POST.get('username', '').strip()
    password = request.POST.get('password', '').strip()

    if not username or not password:
        messages.error(request, "Username va password kiriting.")
        return redirect('login_and_register_page')

    user = authenticate(request, username=username, password=password)

    if user is None:
        messages.error(request, "Login yoki parol xato.")
        return redirect('login_and_register_page')

    login(request, user)
    profile = Profile.objects.filter(user=user).first()

    if profile:
        messages.success(
            request,
            f"Xush kelibsiz, {user.username}. Role: {profile.role}, Code: {profile.unique_code}"
        )
    else:
        messages.warning(request, f"{user.username} uchun profile topilmadi.")

    return redirect('main_page')


@login_required
def logout_from_classroom(request):
    logout(request)
    messages.success(request, "Siz tizimdan chiqdingiz.")
    return redirect('login_and_register_page')


# ─────────────────────────────────────────────────────────────────────────────
# MAIN PAGE
# ─────────────────────────────────────────────────────────────────────────────

@login_required
def main_page(request):
    profile = Profile.objects.filter(user=request.user).first()

    if not profile:
        messages.error(request, "Profile topilmadi.")
        logout(request)
        return redirect('login_and_register_page')

    if request.method == 'POST' and profile.role == 'teacher':
        action     = request.POST.get('action', '').strip()
        request_id = request.POST.get('request_id', '').strip()

        if action in ['accept', 'reject'] and request_id:
            join_request = JoinRequest.objects.filter(
                id=request_id,
                classroom__teacher=request.user
            ).select_related('student', 'classroom').first()

            if not join_request:
                messages.error(request, "Request topilmadi.")
                return redirect('main_page')

            if action == 'accept':
                join_request.status = 'accepted'
                join_request.save()
                created_files = create_missing_files_for_student_in_classroom(
                    student=join_request.student,
                    classroom=join_request.classroom
                )
                messages.success(
                    request,
                    f"{join_request.student.username} uchun request accepted qilindi. "
                    f"Generated files: {created_files}"
                )
            elif action == 'reject':
                join_request.status = 'rejected'
                join_request.save()
                messages.warning(request, f"{join_request.student.username} uchun request rejected qilindi.")

            return redirect('main_page')

    context = {'profile': profile}

    if profile.role == 'teacher':
        context['my_classes'] = Classroom.objects.filter(
            teacher=request.user).order_by('-created_at')
        context['pending_requests_main'] = JoinRequest.objects.filter(
            classroom__teacher=request.user,
            status='pending'
        ).select_related('student', 'student__profile', 'classroom').order_by('-created_at')

    elif profile.role == 'student':
        context['my_requests'] = JoinRequest.objects.filter(
            student=request.user
        ).select_related('classroom', 'classroom__teacher').order_by('-created_at')
        context['joined_classrooms'] = JoinRequest.objects.filter(
            student=request.user,
            status='accepted'
        ).select_related('classroom', 'classroom__teacher').order_by('-created_at')

    return render(request, 'main_page.html', context)


# ─────────────────────────────────────────────────────────────────────────────
# CLASSROOM MANAGEMENT
# ─────────────────────────────────────────────────────────────────────────────

@login_required
def create_classroom(request):
    profile = Profile.objects.filter(user=request.user).first()

    if not profile:
        messages.error(request, "Profile topilmadi.")
        return redirect('login_and_register_page')

    if profile.role != 'teacher':
        messages.error(request, "Faqat teacher class yarata oladi.")
        return redirect('main_page')

    if request.method != 'POST':
        return redirect('main_page')

    name        = request.POST.get('name', '').strip()
    description = request.POST.get('description', '').strip()

    if not name:
        messages.error(request, "Class nomini kiriting.")
        return redirect('main_page')

    classroom = Classroom.objects.create(
        teacher=request.user,
        name=name,
        description=description
    )
    messages.success(request, f"Class yaratildi: {classroom.name}. Code: {classroom.class_code}")
    return redirect('main_page')


@login_required
def join_classroom_request(request):
    profile = Profile.objects.filter(user=request.user).first()

    if not profile:
        messages.error(request, "Profile topilmadi.")
        return redirect('login_and_register_page')

    if profile.role != 'student':
        messages.error(request, "Faqat student classga join request yubora oladi.")
        return redirect('main_page')

    if request.method != 'POST':
        return redirect('main_page')

    class_code = request.POST.get('class_code', '').strip()

    if not class_code:
        messages.error(request, "Class code kiriting.")
        return redirect('main_page')

    classroom = Classroom.objects.filter(class_code=class_code).first()

    if not classroom:
        messages.error(request, "Bunday class topilmadi.")
        return redirect('main_page')

    if classroom.teacher == request.user:
        messages.error(request, "O'zingiz yaratgan classga request yubora olmaysiz.")
        return redirect('main_page')

    existing = JoinRequest.objects.filter(classroom=classroom, student=request.user).first()

    if existing:
        if existing.status == 'pending':
            messages.warning(request, "Siz bu classga allaqachon request yuborgansiz.")
            return redirect('main_page')
        if existing.status == 'accepted':
            messages.warning(request, "Siz bu classga allaqachon qo'shilgansiz.")
            return redirect('main_page')
        if existing.status == 'rejected':
            existing.status = 'pending'
            existing.save()
            messages.success(request, f"{classroom.name} classiga request qayta yuborildi.")
            return redirect('main_page')

    JoinRequest.objects.create(classroom=classroom, student=request.user)
    messages.success(request, f"{classroom.name} classiga join request yuborildi.")
    return redirect('main_page')


# ─────────────────────────────────────────────────────────────────────────────
# CLASSROOM DETAIL
# ─────────────────────────────────────────────────────────────────────────────

@login_required
def classroom_detail_page(request, classroom_id):
    profile = Profile.objects.filter(user=request.user).first()

    if not profile:
        messages.error(request, "Profile topilmadi.")
        logout(request)
        return redirect('login_and_register_page')

    classroom = Classroom.objects.filter(id=classroom_id).select_related('teacher').first()

    if not classroom:
        messages.error(request, "Classroom topilmadi.")
        return redirect('main_page')

    is_teacher = classroom.teacher == request.user
    is_accepted_student = JoinRequest.objects.filter(
        classroom=classroom, student=request.user, status='accepted'
    ).exists()

    if not is_teacher and not is_accepted_student:
        messages.error(request, "Siz bu classroomga kira olmaysiz.")
        return redirect('main_page')

    # ── POST handlers ──────────────────────────────────────────────────────
    if request.method == 'POST':
        action = request.POST.get('action', '').strip()

        # ── Teacher actions ──
        if is_teacher:
            request_id   = request.POST.get('request_id', '').strip()
            student_code = request.POST.get('student_code', '').strip()

            if action in ['accept', 'reject'] and request_id:
                join_request = JoinRequest.objects.filter(
                    id=request_id, classroom=classroom, status='pending'
                ).select_related('student').first()

                if not join_request:
                    messages.error(request, "Pending request topilmadi.")
                    return redirect('classroom_detail_page', classroom_id=classroom.id)

                if action == 'accept':
                    join_request.status = 'accepted'
                    join_request.save()
                    created = create_missing_files_for_student_in_classroom(
                        student=join_request.student, classroom=classroom)
                    messages.success(
                        request,
                        f"{join_request.student.username} accepted. Generated files: {created}"
                    )
                else:
                    join_request.status = 'rejected'
                    join_request.save()
                    messages.warning(request, f"{join_request.student.username} rejected.")

                return redirect('classroom_detail_page', classroom_id=classroom.id)

            if action == 'add_student' and student_code:
                target_profile = Profile.objects.filter(
                    unique_code=student_code).select_related('user').first()

                if not target_profile:
                    messages.error(request, "Bunday student code topilmadi.")
                    return redirect('classroom_detail_page', classroom_id=classroom.id)

                if target_profile.role != 'student':
                    messages.error(request, "Bu code studentga tegishli emas.")
                    return redirect('classroom_detail_page', classroom_id=classroom.id)

                existing = JoinRequest.objects.filter(
                    classroom=classroom, student=target_profile.user).first()

                if existing:
                    if existing.status == 'accepted':
                        messages.warning(
                            request, f"{target_profile.user.username} allaqachon classda bor.")
                        return redirect('classroom_detail_page', classroom_id=classroom.id)
                    existing.status = 'accepted'
                    existing.save()
                else:
                    JoinRequest.objects.create(
                        classroom=classroom, student=target_profile.user, status='accepted')

                created = create_missing_files_for_student_in_classroom(
                    student=target_profile.user, classroom=classroom)
                messages.success(
                    request,
                    f"{target_profile.user.username} direct qo'shildi. Generated files: {created}"
                )
                return redirect('classroom_detail_page', classroom_id=classroom.id)

            if action == 'create_homework':
                title       = request.POST.get('title', '').strip()
                description = request.POST.get('description', '').strip()
                deadline_raw = request.POST.get('deadline', '').strip()
                deadline    = parse_datetime(deadline_raw) if deadline_raw else None

                if not title or not description:
                    messages.error(request, "Homework title va description kiriting.")
                    return redirect('classroom_detail_page', classroom_id=classroom.id)

                homework = Homework.objects.create(
                    classroom=classroom,
                    teacher=request.user,
                    title=title,
                    description=description,
                    deadline=deadline
                )

                attachment_count = 0
                for i in range(1, 11):
                    uploaded_file    = request.FILES.get(f'file_{i}')
                    file_description = request.POST.get(f'file_description_{i}', '').strip()
                    if uploaded_file:
                        HomeworkAttachment.objects.create(
                            homework=homework,
                            file=uploaded_file,
                            description=file_description
                        )
                        attachment_count += 1

                created = create_missing_student_files_for_homework(
                    homework=homework, classroom=classroom)

                messages.success(
                    request,
                    f"Uy ishi yaratildi: {homework.title}. "
                    f"Attachments: {attachment_count}. "
                    f"Student files generated: {created}"
                )
                return redirect('classroom_detail_page', classroom_id=classroom.id)

        # ── Student actions ──
        if is_accepted_student and action == 'submit_homework':
            homework_id = request.POST.get('homework_id', '').strip()
            student_file = StudentHomeworkFile.objects.filter(
                homework_id=homework_id,
                classroom=classroom,
                student=request.user
            ).first()

            if not student_file:
                messages.error(request, "Homework file topilmadi.")
                return redirect('classroom_detail_page', classroom_id=classroom.id)

            if student_file.submitted:
                messages.warning(request, "Siz bu homeworkni allaqachon topshirgansiz.")
                return redirect('classroom_detail_page', classroom_id=classroom.id)

            student_file.submitted    = True
            student_file.submitted_at = timezone.now()
            student_file.save()

            messages.success(request, "Homework muvaffaqiyatli topshirildi!")
            return redirect('classroom_detail_page', classroom_id=classroom.id)

    # ── Build context ──────────────────────────────────────────────────────
    accepted_students = JoinRequest.objects.filter(
        classroom=classroom, status='accepted'
    ).select_related('student', 'student__profile').order_by('-created_at')

    pending_requests = (
        JoinRequest.objects.filter(classroom=classroom, status='pending')
        .select_related('student', 'student__profile')
        .order_by('-created_at')
        if is_teacher else JoinRequest.objects.none()
    )

    homeworks = Homework.objects.filter(
        classroom=classroom
    ).prefetch_related('attachments', 'student_files').order_by('-created_at')

    student_homework_files = StudentHomeworkFile.objects.none()
    if not is_teacher:
        student_homework_files = StudentHomeworkFile.objects.filter(
            classroom=classroom, student=request.user
        ).select_related('homework').order_by('-created_at')

    now = timezone.now()

    context = {
        'profile': profile,
        'classroom': classroom,
        'is_teacher': is_teacher,
        'accepted_students': accepted_students,
        'pending_requests': pending_requests,
        'homeworks': homeworks,
        'student_homework_files': student_homework_files,
        'attachment_range': range(1, 11),
        'now': now,
    }
    return render(request, 'classroom_detail.html', context)


# ─────────────────────────────────────────────────────────────────────────────
# EXIT / KICK
# ─────────────────────────────────────────────────────────────────────────────

@login_required
def exit_or_kick_classroom(request, classroom_id):
    if request.method != 'POST':
        return redirect('main_page')

    profile = Profile.objects.filter(user=request.user).first()
    if not profile:
        messages.error(request, "Profile topilmadi.")
        logout(request)
        return redirect('login_and_register_page')

    classroom = Classroom.objects.filter(id=classroom_id).select_related('teacher').first()
    if not classroom:
        messages.error(request, "Classroom topilmadi.")
        return redirect('main_page')

    action = request.POST.get('action', '').strip()

    if action == 'exit':
        if profile.role != 'student':
            messages.error(request, "Faqat student classdan chiqishi mumkin.")
            return redirect('classroom_detail_page', classroom_id=classroom.id)

        membership = JoinRequest.objects.filter(
            classroom=classroom, student=request.user, status='accepted').first()

        if not membership:
            messages.error(request, "Siz bu classga qo'shilmagansiz.")
            return redirect('classroom_detail_page', classroom_id=classroom.id)

        membership.delete()
        messages.success(request, f"Siz {classroom.name} classidan chiqdingiz.")
        return redirect('main_page')

    if action == 'kick':
        if profile.role != 'teacher' or classroom.teacher != request.user:
            messages.error(request, "Faqat shu classroom o'qituvchisi studentni chiqarib yubora oladi.")
            return redirect('classroom_detail_page', classroom_id=classroom.id)

        student_id = request.POST.get('student_id', '').strip()
        if not student_id:
            messages.error(request, "Student id yuborilmadi.")
            return redirect('classroom_detail_page', classroom_id=classroom.id)

        membership = JoinRequest.objects.filter(
            classroom=classroom, student_id=student_id, status='accepted'
        ).select_related('student').first()

        if not membership:
            messages.error(request, "Bu student shu classda topilmadi.")
            return redirect('classroom_detail_page', classroom_id=classroom.id)

        student_username = membership.student.username
        membership.delete()

        StudentHomeworkFile.objects.filter(
            classroom=classroom, student_id=student_id).delete()

        messages.success(request, f"{student_username} classdan chiqarildi.")
        return redirect('classroom_detail_page', classroom_id=classroom.id)

    messages.error(request, "Noto'g'ri action.")
    return redirect('classroom_detail_page', classroom_id=classroom.id)


# ─────────────────────────────────────────────────────────────────────────────
# BROWSER EDITOR
# ─────────────────────────────────────────────────────────────────────────────

@login_required
def browser_editor_page(request, classroom_id, homework_id):
    classroom = Classroom.objects.filter(id=classroom_id).first()
    homework  = Homework.objects.filter(id=homework_id, classroom=classroom).first()

    if not classroom or not homework:
        messages.error(request, "Homework topilmadi.")
        return redirect('main_page')

    is_teacher = classroom.teacher == request.user
    is_accepted_student = JoinRequest.objects.filter(
        classroom=classroom, student=request.user, status='accepted').exists()

    if not is_teacher and not is_accepted_student:
        messages.error(request, "Ruxsat yo'q.")
        return redirect('main_page')

    student_file = StudentHomeworkFile.objects.filter(
        classroom=classroom, homework=homework, student=request.user).first()

    if not student_file:
        messages.error(request, "Siz uchun homework file topilmadi.")
        return redirect('classroom_detail_page', classroom_id=classroom.id)

    context = {
        'classroom': classroom,
        'homework': homework,
        'student_file': student_file,
        'saved_content': student_file.editor_content or "",
    }
    return render(request, 'browser_editor.html', context)


@require_POST
@login_required
def upload_editor_image(request, classroom_id, homework_id):
    """
    Accepts an image file via POST (multipart), saves it as an EditorImage,
    and returns the served URL so Quill can embed it as a real <img src="...">.
    """
    classroom = Classroom.objects.filter(id=classroom_id).select_related('teacher').first()
    homework  = Homework.objects.filter(id=homework_id, classroom=classroom).first()

    if not classroom or not homework:
        return JsonResponse({'success': False, 'error': 'Homework topilmadi'}, status=404)

    is_accepted_student = JoinRequest.objects.filter(
        classroom=classroom, student=request.user, status='accepted').exists()

    if not is_accepted_student:
        return JsonResponse({'success': False, 'error': 'Ruxsat yo\'q'}, status=403)

    image_file = request.FILES.get('image')
    if not image_file:
        return JsonResponse({'success': False, 'error': 'Fayl yuborilmadi'}, status=400)

    # Validate it's actually an image
    allowed_types = ('image/jpeg', 'image/png', 'image/gif', 'image/webp')
    if image_file.content_type not in allowed_types:
        return JsonResponse({'success': False, 'error': 'Faqat rasm fayllari qabul qilinadi'}, status=400)

    # Limit size to 10 MB
    if image_file.size > 10 * 1024 * 1024:
        return JsonResponse({'success': False, 'error': 'Rasm 10MB dan kichik bo\'lishi kerak'}, status=400)

    editor_img = EditorImage.objects.create(
        homework=homework,
        student=request.user,
        file=image_file,
    )

    return JsonResponse({
        'success': True,
        'url': editor_img.file.url,
    })


@require_POST
@login_required
def save_browser_editor(request, classroom_id, homework_id):
    """
    Saves editor HTML to editor_content field AND regenerates the physical
    .docx file with the correct naming convention so downloads always reflect
    the latest content — including embedded images.
    """
    classroom = Classroom.objects.filter(id=classroom_id).select_related('teacher').first()
    homework  = Homework.objects.filter(
        id=homework_id, classroom=classroom
    ).prefetch_related('attachments').first()

    if not classroom or not homework:
        return JsonResponse({'success': False, 'error': 'Homework topilmadi'}, status=404)

    student_file = StudentHomeworkFile.objects.filter(
        classroom=classroom, homework=homework, student=request.user).first()

    if not student_file:
        return JsonResponse({'success': False, 'error': 'File topilmadi'}, status=404)

    data    = json.loads(request.body)
    content = data.get('content', '')

    # 1. Save HTML content to DB
    student_file.editor_content = content

    # 2. Regenerate the .docx with the correct filename and the new content
    _refresh_docx_file(student_file)

    # 3. Persist both changes
    student_file.save()

    return JsonResponse({'success': True})